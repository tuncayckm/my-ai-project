# pip install langchain
from pydantic_settings import BaseSettings, SettingsConfigDict
from transformers import pipeline
from pydantic import Field
from typing import Optional
from your_mlops_tool import get_active_model_name_from_registry # Örnek bir kod
from langchain_text_splitters import RecursiveCharacterTextSplitter, NLTKTextSplitter
from langchain.text_splitters import SemanticChunker # Bu kütüphane için
from langchain_community.embeddings import HuggingFaceEmbeddings # Veya OpenAIEmbeddings
from langchain_core.documents import Document
import os
from embedding import save_embedding_to_store
import whisper
import redis.asyncio as aioredis
import time
import docx
import fitz
import asyncio
import logging
import aiohttp
import aiofiles
import re
from config import app_settings
from bs4 import BeautifulSoup
from pydub import AudioSegment
import speech_recognition as sr
from transformers import pipeline, AutoTokenizer
from embedding import save_memory
from typing import List, Optional
from functools import lru_cache
from pathlib import Path
import tempfile
import validators
from prometheus_client import Counter, Histogram
from tenacity import retry, stop_after_attempt, wait_fixed
from langchain_text_splitters import RecursiveCharacterTextSplitter, NLTKTextSplitter
import nltk
from nltk.tokenize import sent_tokenize
from gettext import gettext as _
import redis
import jwt
import hashlib
# --- i18n / gettext ayarları ---
import gettext
import locale

LOCALE_DIR = os.path.join(os.path.dirname(__file__), "locales")
DEFAULT_LANG = os.getenv("APP_LANGUAGE", "en")


class LearningService:
    """
    Modülün tüm ana iş akışını yöneten merkezi servis sınıfı.
    Bağımlılıkları (settings, rate_limiter, redis_client) enjeksiyonla alır.
    """
    def __init__(self, rate_limiter, redis_client):
        self.settings = app_settings
        self.rate_limiter = rate_limiter
        self.redis_client = redis_client
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'
        }

    async def learn_from_text(self, user_id: str, text: str, source: str = "general"):
        """
        Belirli bir kullanıcı için metinden öğrenir ve önbelleğe alır.
        Bu versiyon, metni anlamsal olarak ilgili parçalara ayırmak için SemanticChunker kullanır.
        """
        start = time.time()
        
        text_hash = hashlib.sha256(text.encode('utf-8')).hexdigest()
        cache_key = f"processed_text:{user_id}:{text_hash}"
        
        async_redis_client = aioredis_from_url(f"redis://{self.settings.redis_host}")

        try:
            if await async_redis_client.exists(cache_key):
                log_i18n("info", f"Metin önbellekte bulundu. İşlem atlanıyor. User: {user_id}, Hash: {text_hash}")
                return

            text_splitter = get_text_splitter(splitter_type="semantic")
            docs = [Document(page_content=text)]
            chunks = text_splitter.split_documents(docs)
            chunks_content = [chunk.page_content for chunk in chunks]

            log_i18n("info", f"Metin {len(chunks_content)} parçaya ayrıldı. User: {user_id}")

            tasks = [save_memory(user_id, chunk_content, source) for chunk_content in chunks_content]
            await asyncio.gather(*tasks)

            await async_redis_client.setex(cache_key, self.settings.text_cache_ttl_sec, "1")
            log_i18n("info", f"Metin başarıyla işlendi ve önbelleğe alındı. User: {user_id}")

            LEARN_COUNTER.labels(source=source).inc()
        except Exception as e:
            prometheus_log_exception(source)
            LEARN_ERROR_COUNTER.labels(source=source).inc()
            log_i18n("error", f"Metinden öğrenme hatası: {e}")
        finally:
            PROCESSING_TIME.labels(source=source).observe(time.time() - start)

    async def learn_from_pdf(self, user_id: str, file_path: str):
        log_i18n("info", f"PDF dosyasından öğrenme başlatılıyor: {file_path}")
        start = time.time()
        
        if not validate_file_path(file_path, allowed_suffixes=[".pdf"]):
            log_i18n("error", "Geçersiz PDF dosyası veya dosya yolu.")
            return

        try:
            doc = fitz.open(file_path)
            full_text = ""
            extracted_images = []

            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                full_text += page.get_text() + "\n"
                
                # Görselleri çıkar
                for img_info in doc.get_page_images(page_num):
                    xref = img_info[0]
                    pix = fitz.Pixmap(doc, xref)
                    
                    if pix.n - pix.alpha < 4:  # RGB veya Gri tonlama
                        img_path = f"{file_path}_{page_num}_{xref}.png"
                        pix.save(img_path)
                        extracted_images.append(img_path)

            doc.close()

            # Görsel altyazılarını oluştur
            if extracted_images:
                log_i18n("info", "Görseller için altyazılar oluşturuluyor...")
                image_captions = await asyncio.gather(
                    *[generate_image_caption(img_path) for img_path in extracted_images]
                )
                
                # Oluşturulan altyazıları ana metne ekle
                combined_text = full_text + "\n" + "\n".join(image_captions)
                await self.learn_from_multimodal_text(user_id, combined_text, source="pdf")
            else:
                await self.learn_from_text(user_id, full_text, source="pdf")
            
            log_i18n("info", "PDF işleme tamamlandı.")
            return True
        except Exception as e:
            prometheus_log_exception("pdf")
            log_i18n("error", f"PDF işleme hatası: {e}")
            return False
        finally:
            PROCESSING_TIME.labels(source="pdf").observe(time.time() - start)
            # Geçici dosyaları temizle
            for img_path in extracted_images:
                os.remove(img_path)

    async def learn_from_docx(self, user_id: str, file_path: str):
        log_i18n("info", f"DOCX dosyasından öğrenme başlatılıyor: {file_path}")
        start = time.time()

        if not validate_file_path(file_path, allowed_suffixes=[".docx"]):
            log_i18n("error", "Geçersiz DOCX dosyası veya dosya yolu.")
            return

        try:
            doc = DocxDocument(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            await self.learn_from_text(user_id, full_text, source="docx")
            log_i18n("info", "DOCX işleme tamamlandı.")
            return True
        except Exception as e:
            prometheus_log_exception("docx")
            log_i18n("error", f"DOCX işleme hatası: {e}")
            return False
        finally:
            PROCESSING_TIME.labels(source="docx").observe(time.time() - start)

    async def learn_from_url(self, user_id: str, url: str):
        log_i18n("info", f"URL'den öğrenme başlatılıyor: {url}")
        start = time.time()
        
        async with aiohttp.ClientSession(headers=self.headers) as session:
            try:
                async with session.get(url) as response:
                    if response.status == 200:
                        html_content = await response.text()
                        soup = BeautifulSoup(html_content, 'html.parser')
                        
                        # script ve style etiketlerini temizle
                        for script_or_style in soup(['script', 'style']):
                            script_or_style.extract()
                            
                        text_content = soup.get_text()
                        cleaned_text = " ".join(text_content.split())
                        
                        # HTML temizleme ve içerik filtrelemesini uygula
                        cleaned_text = clean_html(cleaned_text)
                        
                        if advanced_content_filter(cleaned_text):
                            await self.learn_from_text(user_id, cleaned_text, source="url")
                            log_i18n("info", "URL işleme tamamlandı.")
                            return True
                        else:
                            log_i18n("warning", "URL içeriği güvenlik filtresinden geçemedi.")
                            return False
                    else:
                        log_i18n("error", f"URL'e erişilemedi: {url} (Status: {response.status})")
                        return False
            except Exception as e:
                prometheus_log_exception("url")
                log_i18n("error", f"URL işleme hatası: {e}")
                return False
            finally:
                PROCESSING_TIME.labels(source="url").observe(time.time() - start)


def extract_images_from_pdf(pdf_path: str):
    """
    PDF dosyasından görselleri çıkarır ve geçici bir klasöre kaydeder.
    Görsel yollarının bir listesini döndürür.
    """
    image_paths = []
    try:
        doc = fitz.open(pdf_path)
        temp_dir = "temp_images"
        os.makedirs(temp_dir, exist_ok=True)
        
        for page_index in range(len(doc)):
            for img_index, img in enumerate(doc.get_page_images(page_index)):
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                if pix.n - pix.alpha > 3:  # CMYK renk alanını RGB'ye dönüştür
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                
                image_path = os.path.join(temp_dir, f"img_{uuid.uuid4()}.png")
                pix.save(image_path)
                image_paths.append(image_path)
                pix = None  # Bellek sızıntısını önler
        
        doc.close()
    except Exception as e:
        prometheus_log_exception(source)
        log_i18n(f"PDF'den görsel çıkarma hatası: {e}")
        
    return image_paths

async def generate_image_caption(image_path: str) -> str:
    """
    Bir multi-modal LLM çağrısı ile görselin içeriğini özetler.
    Bu, Google Gemini API'sinin kullanıldığı bir örnektir.
    """
    log_i18n(f"Görsel için altyazı oluşturuluyor: {image_path}")
    
    # 1. API Anahtarınızı Ortam Değişkeninden Alın
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
    if not GEMINI_API_KEY:
        log_i18n("error", "GEMINI_API_KEY ortam değişkeni tanımlı değil.")
        return "[Görsel Açıklaması: API anahtarı eksik olduğu için açıklama oluşturulamadı.]"

    # 2. Gemini Kütüphanesini Kurun ve İçe Aktarın
    # pip install google-generativeai
    try:
        import google.generativeai as genai
    except ImportError:
        log_i18n("error", "google-generativeai kütüphanesi kurulu değil.")
        return "[Görsel Açıklaması: Gerekli kütüphane eksik.]"
    
    genai.configure(api_key=GEMINI_API_KEY)

    # 3. Model ve İstek Hazırlığı
    model = genai.GenerativeModel('gemini-pro-vision')

    with open(image_path, "rb") as image_file:
        image_data = image_file.read()

    # 4. Asenkron API Çağrısı
    try:
        response = await model.generate_content_async(
            ["Bu resimde ne görüyorsun? Detaylı bir açıklama yap.", image_data],
            stream=True
        )
        caption_parts = []
        async for chunk in response:
            caption_parts.append(chunk.text)

        full_caption = "".join(caption_parts)
        log_i18n(f"Görsel altyazısı başarıyla oluşturuldu.")
        return f"[Görsel Açıklaması: {full_caption}]"

    except Exception as e:
        prometheus_log_exception("multimodal_api")
        log_i18n("error", f"Gemini API çağrısı hatası: {e}")
        return f"[Görsel Açıklaması: Oluşturma hatası - {e}]"

def get_active_model_name(model_type: str):
    if model_type == "summarizer":
        return settings.summary_model_name
    elif model_type == "tokenizer":
        return settings.tokenizer_model_name
    return None

try:
    language = gettext.translation('messages', localedir=LOCALE_DIR, languages=[DEFAULT_LANG])
    language.install()
    _ = language.gettext
except FileNotFoundError:
    # Fallback: İngilizce
    gettext.install('messages')
    _ = gettext.gettext

# --- NLTK punkt tokenizer kurulumu ---
# nltk.download('punkt')  # Ortamda sadece 1 kere veya Dockerfile'da yapılmalı

# --- Logger ayarları ---
logging.basicConfig(format='%(asctime)s - %(levelname)s - %(message)s', level=logging.INFO)
logger = logging.getLogger(__name__)


def log_i18n(level: str, message: str, **kwargs):
    """
    Çok dilli log mesajları için yardımcı fonksiyon.
    Örnek: log_i18n("warning", "Dosya bulunamadı: %(file_path)s", file_path="abc.pdf")
    """
    localized_msg = _(message) % kwargs
    getattr(logger, level)(localized_msg)

# --- Config ---


if settings.environment == "production" and not settings.jwt_secret:
    raise ValueError("Üretim ortamında JWT_SECRET zorunludur.")

# --- Redis Bağlantısı ---
try:
    redis_client = redis.Redis(host=settings.redis_host, port=REDIS_PORT, db=REDIS_DB, decode_responses=True)
    redis_client.ping()
except redis.ConnectionError as e:
    log_i18n(f"Redis bağlantı hatası: {e}")
    redis_client = None

# --- Rate Limiter Sınıfı ---

@retry(stop=stop_after_attempt(3), wait=wait_fixed(2))
def connect_to_redis(host, port, db):
    return redis.Redis(host=host, port=port, db=db, decode_responses=True)

class RedisRateLimiter:
    def __init__(self, key_prefix, rate, per):
        self.key_prefix = key_prefix
        self.rate = rate
        self.per = per
        try:
            self.redis = connect_to_redis(
                host=settings.redis_host,
                port=settings.redis_port,
                db=settings.redis_db
            )
        except Exception as e:
            prometheus_log_exception(source)
            log_i18n(f"Redis bağlantı hatası, yeniden deneme mekanizması devrede: {e}")
            # Hata durumunda fail-safe olarak None atayın veya başka bir işlem yapın.
            self.redis = None

    def _key(self, user_id: str) -> str:
        return f"{self.key_prefix}:{user_id}"

    def allow(self, user_id: str) -> bool:
        if redis_client is None:
            log_i18n("Redis yok, rate limiting devre dışı.")
            return True

        key = self._key(user_id)
        now = int(time.time())
        try:
            with redis_client.pipeline() as pipe:
                pipe.zremrangebyscore(key, 0, now - self.per)
                pipe.zadd(key, {str(now): now})
                pipe.zcard(key)
                pipe.expire(key, self.per + 1)
                results = pipe.execute()
            current_count = results[2]
            if current_count > self.rate:
                logger.debug(f"Rate limit aşıldı: user_id={user_id}, count={current_count}")
                return False
            return True
        except redis.RedisError as e:
            log_i18n(f"Redis rate limiter hatası: {e}")
            return True

rate_limiter = RedisRateLimiter(key_prefix="learning_rate", rate=settings.rate_limit, per=settings.rate_limit_period)

# --- Prometheus Metrikleri ---
REQUEST_COUNTER = Counter("learning_requests_total", "Toplam öğrenme isteği sayısı", ["source"])
ERROR_COUNTER = Counter("learning_errors_total", "Toplam öğrenme hatası sayısı", ["source"])
PROCESSING_TIME = Histogram("learning_processing_seconds", "Öğrenme süresi saniye", ["source"])

def prometheus_log_exception(source: str):
    ERROR_COUNTER.labels(source=source).inc()

# --- Celery ---
celery_app = celery.Celery('learning_tasks', broker=settings.celery_broker_url) # settings'e eklemeniz gerekir

# --- Tokenizer ve Özetleyici ---


@lru_cache(maxsize=1)
def get_summarizer():
    model_name = get_active_model_name_from_registry("summarizer") or "sshleifer/distilbart-cnn-12-6"
    try:
        return pipeline("summarization", model=model_name)
    except Exception as e:
        prometheus_log_exception("summarizer")
        log_i18n("warning", "Özetleyici yüklenemedi: %(err)s", err=str(e))
        return pipeline("summarization", model="facebook/bart-large-cnn")

@lru_cache(maxsize=1)
def get_tokenizer():
    active_model = get_active_model_name("tokenizer")
    if active_model:
        try:
            return AutoTokenizer.from_pretrained(active_model)
        except Exception as e:
            prometheus_log_exception(source)
            log_i18n(f"Tokenizer yüklenemedi: {e}")
    return None

@lru_cache(maxsize=1)
def get_text_splitter(splitter_type: str = "recursive", chunk_size: int = 500, chunk_overlap: int = 50):
    if splitter_type == "semantic":
        embeddings = HuggingFaceEmbeddings(model_name=settings.embedding_model_name)
        return SemanticChunker(embeddings)
    elif splitter_type == "nltk":
        return NLTKTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap, length_function=len)


# --- Gelişmiş İçerik Filtresi ---
@lru_cache(maxsize=1)
def get_content_classifier():
    """Zararlı içerik tespiti için NLP sınıflandırıcısını yükler."""
    try:
        # Bu model örnektir, daha uygun bir modelle değiştirilebilir
        return pipeline("text-classification", model="unitary/toxic-bert", framework="pt")
    except Exception as e:
        log_i18n("warning", f"Zararlı içerik sınıflandırıcısı yüklenemedi: {e}")
        return None

def advanced_content_filter(text: str) -> bool:
    classifier = get_content_classifier()
    if not classifier:
        # Fallback olarak anahtar kelime kontrolü
        spam_keywords = ['spamword1', 'malware', 'phishing', 'hack', 'attack']
        text_lower = text.lower()
        for kw in spam_lower:
            if kw in text_lower:
                log_i18n("warning", "Fallback olarak anahtar kelime ile zararlı içerik tespit edildi.")
                return False
        return True

    result = classifier(text)
    # Örnek bir mantık, modelin çıktısına göre güncellenmeli
    # 'toxic', 'severe_toxic', 'obscene', 'threat', 'insult', 'identity_hate'
    for label in result:
        if label['label'] in ['toxic', 'threat'] and label['score'] > 0.8:
            log_i18n("warning", f"NLP modeli ile zararlı içerik tespit edildi: {label}")
            return False
    return True

# --- HTML Temizleme ---
def clean_html(raw_html: str) -> str:
    soup = BeautifulSoup(raw_html, "html.parser")
    for tag in soup(["script", "style", "header", "footer", "nav", "aside", "form"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    clean_text = "\n".join(lines)
    # TODO: XSS ve virüs taraması entegre edilmeli
    return clean_text

# --- Dosya Doğrulama ---
def validate_file_path(file_path: str, allowed_suffixes: Optional[List[str]] = None, max_size_mb: Optional[int] = None) -> bool:
    allowed_suffixes = allowed_suffixes or [".pdf", ".docx", ".txt", ".epub"]
    max_size_mb = max_size_mb or settings.max_file_size_mb
    p = Path(file_path)
    if not p.exists():
        log_i18n(f"Dosya bulunamadı: {file_path}")
        return False
    if p.stat().st_size > max_size_mb * 1024 * 1024:
        log_i18n(f"Dosya çok büyük: {file_path}")
        return False
    if p.suffix.lower() not in allowed_suffixes:
        log_i18n(f"Desteklenmeyen dosya türü: {file_path}")
        return False
    return True

# --- Metin Önbellekleme ---
def is_text_cached(user_id: str, text_hash: str) -> bool:
    if redis_client is None:
        return False
    key = f"text_cache:{user_id}:{text_hash}"
    try:
        return redis_client.exists(key) == 1
    except redis.RedisError as e:
        log_i18n("error", "Redis önbellek kontrol hatası: %(error)s", error=str(e))
        return False

def cache_text(user_id: str, text_hash: str, ttl_sec: int = TEXT_CACHE_TTL_SEC):
    if redis_client is None:
        return
    key = f"text_cache:{user_id}:{text_hash}"
    try:
        redis_client.setex(key, ttl_sec, "1")
    except redis.RedisError as e:
        log_i18n(f"Redis önbellek set hatası: {e}")

def hash_text(text: str) -> str:
    return hashlib.sha256(text.encode('utf-8')).hexdigest()

# --- JWT Yetkilendirme ---
def check_user_authorization(user_jwt_token: str, resource_id: str) -> bool:
    try:
        payload = jwt.decode(user_jwt_token, settings.jwt_secret, algorithms=[settings.jwt_algorithm])
        user_id = payload.get("user_id")
        scopes = payload.get("scopes", [])
        if not user_id or "read" not in scopes:
            log_i18n("JWT token eksikliği.")
            return False
        # Redis'in SET yapısı ile daha güvenli kara liste kontrolü
        if redis_client and redis_client.sismember("jwt_blacklist", user_jwt_token):
            log_i18n("warning", "JWT token kara listede.")
            return False
        # TODO: resource_id bazlı gerçek erişim kontrolü (veritabanı veya servis entegrasyonu)
        return True
    except jwt.ExpiredSignatureError:
        log_i18n("JWT token süresi dolmuş.")
        return False
    except jwt.InvalidTokenError:
        log_i18n("Geçersiz JWT token.")
        return False
    except Exception as e:
        prometheus_log_exception(source)
        log_i18n(f"JWT doğrulama hatası: {e}")
        return False

# --- Asenkron Chunk Kaydetme ---
async def async_learn_text_chunks(user_id: str, chunks: List[str]):
    tasks = []
    for chunk in chunks:
        if not rate_limiter.allow(user_id):
            log_i18n(f"Rate limit aşıldı, işlem bekletiliyor. user_id={user_id}")
            await asyncio.sleep(0.2)
        try:
            tasks.append(save_memory(user_id, chunk, source="text"))
        except Exception as e:
            prometheus_log_exception("general")
            log_i18n(f"save_memory çağrısında hata: {e}")
    results = await asyncio.gather(*tasks, return_exceptions=True)
    for i, res in enumerate(results):
        if isinstance(res, Exception):
            log_i18n(f"Chunk {i} kaydedilirken hata: {res}")
        else:
            log_i18n(f"Chunk {i} başarıyla kaydedildi.")

# --- Celery Task: Özetle ve Kaydet ---
@lru_cache(maxsize=1)
def get_sentiment_analyzer():
    log_i18n("Loading sentiment analysis model...")
    return pipeline("sentiment-analysis")

@shared_task
def celery_summarize_and_save(user_id: str, text_or_file_path: str, file_type: str = "text"):
    """
    Metin, PDF veya diğer dosyalardan öğrenme işlemini gerçekleştirir.
    """
    if file_type == "pdf":
        try:
            # Görsel çıkarma ve işleme
            image_paths = extract_images_from_pdf(text_or_file_path)
            for path in image_paths:
                caption = generate_caption(path)
                # save_memory fonksiyonunuzu güncelleyerek hem metin hem de resim verilerini saklayın.
                save_memory(user_id, caption, source="image")
                os.remove(path) # Geçici dosyaları silin
        except Exception as e:
            prometheus_log_exception(source)
            log_i18n(f"Görsel işleme hatası: {e}")

    # Metin işleme kısmı (mevcut mantığınızı burada devam ettirin)
    if not is_text_cached(user_id, text_or_file_path):
        text_content = get_text_from_source(text_or_file_path, file_type)
        if text_content:
            summarizer = get_summarizer()
            text_splitter = get_text_splitter("recursive")
            chunks = text_splitter.split_text(text_content)

            for chunk in chunks:
                summary_output = summarizer(chunk, max_length=150, min_length=30)
                summary = summary_output[0]['summary_text']
                save_memory(user_id, summary, source="text")
                
            cache_text(user_id, hash_text(text_content))

# --- Ses Tanıma Celery Task ---
@celery_app.task(bind=True, max_retries=3, default_retry_delay=10)
def celery_audio_to_text_and_save(self, user_id: str, audio_path: str):
    recognizer = sr.Recognizer()
    audio_file = Path(audio_path)
    if not audio_file.exists():
        log_i18n(f"Ses dosyası bulunamadı (celery): {audio_path}")
        return
    try:
        sound = AudioSegment.from_file(audio_path)
        chunk_length_ms = 60 * 1000
        chunks = [sound[i:i+chunk_length_ms] for i in range(0, len(sound), chunk_length_ms)]

        full_text = ""
        for i, chunk in enumerate(chunks):
            with tempfile.NamedTemporaryFile(suffix=f"_{i}.wav", delete=True) as tmpfile:
                chunk.export(tmpfile.name, format="wav")
                with sr.AudioFile(tmpfile.name) as source:
                    audio = recognizer.record(source)
                try:
                    text = recognizer.recognize_google(audio, language="tr-TR")
                    full_text += text + " "
                    log_i18n(f"Chunk {i} başarıyla dönüştürüldü.")
                except sr.UnknownValueError:
                    log_i18n(f"Chunk {i} için konuşma anlaşılamadı.")
                except sr.RequestError as e:
                    log_i18n(f"Google Speech API hatası: {e}")
                    raise self.retry(exc=e)

        if full_text:
            save_memory(user_id, full_text, source="audio")
            log_i18n(f"Ses dosyasından metin başarıyla kaydedildi.")
        else:
            log_i18n("Ses dosyasından metin çıkmadı.")
    except Exception as e:
        prometheus_log_exception("audio")
        log_i18n(f"Ses dosyası işleme hatası: {e}")
        raise self.retry(exc=e)

def transcribe_audio(audio_path: str, engine: str = "google") -> str:
    if engine == "whisper":
        model = whisper.load_model("base")
        result = model.transcribe(audio_path)
        return result["text"]
    else:
        recognizer = sr.Recognizer()
        with sr.AudioFile(audio_path) as source:
            audio = recognizer.record(source)
        return recognizer.recognize_google(audio, language="tr-TR")


# --- Basit Test Fonksiyonları ---
def test_validate_file_path():
    assert validate_file_path(__file__, allowed_suffixes=[".py"]) == True
    assert validate_file_path("olmayan_dosya.txt") == False
    print("validate_file_path testi başarılı")

def test_rate_limiter():
    limiter = RedisRateLimiter("test_key", rate=2, per=1)
    user = "user123"
    allowed_1 = limiter.allow(user)
    allowed_2 = limiter.allow(user)
    allowed_3 = limiter.allow(user)
    assert allowed_1 == True
    assert allowed_2 == True
    assert allowed_3 == False
    print("rate_limiter testi başarılı")

def test_jwt_auth():
    token = jwt.encode({"user_id": "testuser"}, JWT_SECRET, algorithm=JWT_ALGORITHM)
    assert check_user_authorization(token, "resource1") == True
    assert check_user_authorization("gecersiztoken", "resource1") == False
    print("jwt_auth testi başarılı")

if __name__ == "__main__":
    test_validate_file_path()
    test_rate_limiter()
    test_jwt_auth()
    log_i18n("Tüm testler başarıyla tamamlandı.")
